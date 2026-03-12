"""
APA Validator AI
Sistema de Validación de Citas y Referencias APA mediante Inteligencia Artificial

Copyright (c) 2026 Michael González
All rights reserved.

Este software es propiedad de Michael González.
Queda prohibida la reproducción, distribución, modificación
o uso del código sin autorización expresa del autor.
"""
import streamlit as st
import docx
import io
import os
from openai import OpenAI
from dotenv import load_dotenv

# Librerías para RAG con LangChain
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS

# 1. CARGA DE CONFIGURACIÓN
load_dotenv()
API_KEY = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")

# --- CONFIGURACIÓN DE UI ---
st.set_page_config(page_title="Agente APA 7 - Biblioteca con RAG", page_icon="📚", layout="wide")

# --- MOTOR DE CONOCIMIENTO (RAG) ---
@st.cache_resource
def inicializar_conocimiento(ruta_pdf):
    """Carga el PDF, lo fragmenta y crea un índice de búsqueda (Vector Store)."""
    try:
        loader = PyPDFLoader(ruta_pdf)
        paginas = loader.load()
        
        # Dividir el manual en fragmentos lógicos de 1000 caracteres
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        docs = text_splitter.split_documents(paginas)
        
        # Crear la base de datos vectorial (FAISS) usando Embeddings de OpenAI
        embeddings = OpenAIEmbeddings(openai_api_key=API_KEY)
        vector_store = FAISS.from_documents(docs, embeddings)
        return vector_store
    except Exception as e:
        st.error(f"Error al cargar la base de conocimiento: {e}")
        return None

# --- LÓGICA DE NEGOCIO ---
def extraer_secciones(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    cuerpo, refs = [], []
    en_refs = False
    marcas = ["referencias", "bibliografía", "lista de referencias", "obras citadas"]
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        if t.lower().replace(":", "") in marcas:
            en_refs = True
            continue
        if en_refs: refs.append(t)
        else: cuerpo.append(t)
    return "\n".join(cuerpo), "\n".join(refs)

def analizar_con_rag(cuerpo, refs, vector_store):
    """Consulta el manual PDF y genera el reporte basado en evidencia."""
    # 1. Buscar en el manual las reglas relevantes para el texto del alumno
    consulta = f"Reglas de citas y referencias para: {cuerpo[:500]}... {refs[:500]}"
    docs_relevantes = vector_store.similarity_search(consulta, k=4)
    contexto_manual = "\n\n".join([d.page_content for d in docs_relevantes])
    
    client = OpenAI(api_key=API_KEY)
    prompt_sistema = (
        "Eres un bibliotecario experto en APA 7. Tu tarea es generar un REPORTE DE FEEDBACK.\n"
        "IMPORTANTE: Utiliza el CONTEXTO DEL MANUAL proporcionado para justificar tus correcciones.\n"
        "1. COHERENCIA: Citas en texto vs Referencias.\n"
        "2. FORMATO APA 7: Basado estrictamente en el manual adjunto.\n"
        "3. CORRECCIONES: Cita la regla del manual si es posible."
    )
    
    prompt_usuario = (
        f"--- CONTEXTO DEL MANUAL OFICIAL ---\n{contexto_manual}\n\n"
        f"--- TRABAJO DEL ALUMNO ---\nCUERPO:\n{cuerpo}\n\nREFS:\n{refs}"
    )
    
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "system", "content": prompt_sistema},
                  {"role": "user", "content": prompt_usuario}]
    )
    return response.choices[0].message.content

# --- INTERFAZ PRINCIPAL ---
st.title("📚 Agente APA 7 con Base de Conocimiento (RAG)")
st.sidebar.header("Estado del Sistema")

if not API_KEY:
    st.error("❌ Falta OPENAI_API_KEY.")
    st.stop()

# Inicializar base de conocimiento (el archivo debe llamarse manual_apa7.pdf)
ruta_manual = "manual_apa7.pdf"
if os.path.exists(ruta_manual):
    vector_db = inicializar_conocimiento(ruta_manual)
    st.sidebar.success("✅ Manual APA 7 cargado y activo.")
else:
    st.sidebar.warning("⚠️ No se encontró 'manual_apa7.pdf'. El agente usará su conocimiento general.")
    vector_db = None

archivo = st.file_uploader("Sube el trabajo del alumno (.docx)", type=["docx"])

if archivo:
    if st.button("🚀 Iniciar Análisis con Base de Conocimiento"):
        with st.spinner("Consultando el manual y analizando el documento..."):
            cuerpo, refs = extraer_secciones(archivo.read())
            
            if vector_db:
                resultado = analizar_con_rag(cuerpo, refs, vector_db)
            else:
                # Fallback si no hay RAG
                resultado = "Error: No se pudo acceder a la base de conocimiento."
            
            st.success("¡Análisis Finalizado!")
            st.markdown(resultado)
            
            # Generar Word descargable
            doc_out = docx.Document()
            doc_out.add_heading(f"Reporte APA (Basado en Manual) - {archivo.name}", 0)
            for line in resultado.split('\n'):
                doc_out.add_paragraph(line)
            
            buffer = io.BytesIO()
            doc_out.save(buffer)
            st.download_button("📥 Descargar Reporte Profesional", buffer.getvalue(), 
                               file_name=f"Feedback_RAG_{archivo.name}", 
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
