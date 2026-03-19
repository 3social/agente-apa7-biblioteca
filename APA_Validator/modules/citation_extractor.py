"""
Extractor de secciones de documentos APA 7.

Modo básico  (extractor_completo=False): extrae cuerpo y referencias.
             Comportamiento original — cero riesgo de regresión.

Modo completo (extractor_completo=True): detecta todas las secciones
             mediante una máquina de estados que recorre el documento
             en orden lineal:
             Portada → Abstract → Cuerpo (H1-H5) → Referencias → Apéndices
             + detección lateral de Tablas y Figuras.
"""

import docx
import io
from .schemas import DocumentoAPA, Encabezado


# ── Marcas de sección ─────────────────────────────────────────────────────────
# Normalizadas a minúsculas sin dos puntos ni punto final.
# Se aceptan variantes en español e inglés para cubrir documentos bilingües.

_MARCAS_ABSTRACT = {
    "resumen", "abstract", "resumen ejecutivo", "executive summary",
}
_MARCAS_REFS = {
    "referencias", "bibliografía", "lista de referencias",
    "obras citadas", "references", "bibliography",
}
_MARCAS_APENDICE = {
    "apéndice", "apendice", "appendix",
    "apéndices", "apendices", "appendices",
}

# Nombres de estilos de encabezado en Word (español e inglés).
# Se busca con .lower() para ignorar mayúsculas del tema de Word.
_ESTILOS_ENCABEZADO: dict[str, int] = {
    "heading 1": 1, "título 1": 1, "encabezado 1": 1,
    "heading 2": 2, "título 2": 2, "encabezado 2": 2,
    "heading 3": 3, "título 3": 3, "encabezado 3": 3,
    "heading 4": 4, "título 4": 4, "encabezado 4": 4,
    "heading 5": 5, "título 5": 5, "encabezado 5": 5,
}


# ── Utilidades internas ───────────────────────────────────────────────────────

def _normalizar(texto: str) -> str:
    """Limpia el texto para comparar con las marcas de sección."""
    return texto.strip().lower().rstrip(":.").strip()


def _nivel_encabezado(parrafo) -> int | None:
    """
    Retorna el nivel (1–5) si el párrafo tiene un estilo de encabezado Word.
    Retorna None si es un párrafo de cuerpo normal.
    """
    estilo = parrafo.style.name.lower()
    return _ESTILOS_ENCABEZADO.get(estilo)


def _es_titulo_tabla(texto: str) -> bool:
    t = texto.lower()
    return t.startswith("tabla ") or t.startswith("table ")


def _es_titulo_figura(texto: str) -> bool:
    t = texto.lower()
    return t.startswith("figura ") or t.startswith("figure ")


# ── Extractores ───────────────────────────────────────────────────────────────

def _extraer_basico(doc) -> DocumentoAPA:
    """
    Modo legacy: extrae solo cuerpo y referencias.
    Comportamiento idéntico al extractor original — seguro para rollback.
    """
    cuerpo, refs = [], []
    en_refs = False

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if _normalizar(t) in _MARCAS_REFS:
            en_refs = True
            continue
        if en_refs:
            refs.append(t)
        else:
            cuerpo.append(t)

    return DocumentoAPA(
        cuerpo="\n".join(cuerpo),
        referencias="\n".join(refs),
    )


def _extraer_completo(doc) -> DocumentoAPA:
    """
    Modo completo: detecta todas las secciones APA 7 con máquina de estados.

    Estados:
      PORTADA     → párrafos iniciales antes del resumen.
      ABSTRACT    → sección resumen/abstract.
      CUERPO      → desarrollo del trabajo con encabezados H1-H5.
      REFERENCIAS → lista de referencias.
      APENDICES   → apéndices al final del documento.

    Detección lateral (no cambia estado):
      - Títulos de tablas (Tabla N, Table N).
      - Títulos de figuras (Figura N, Figure N).
    """
    PORTADA     = "portada"
    ABSTRACT    = "abstract"
    CUERPO      = "cuerpo"
    REFERENCIAS = "referencias"
    APENDICES   = "apendices"

    estado = PORTADA

    portada:     list[str]        = []
    abstract:    list[str]        = []
    cuerpo:      list[str]        = []
    refs:        list[str]        = []
    apendices:   list[str]        = []
    encabezados: list[Encabezado] = []
    tablas:      list[str]        = []
    figuras:     list[str]        = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue

        norm  = _normalizar(t)
        nivel = _nivel_encabezado(p)

        # ── Transiciones de estado ────────────────────────────────────────────

        if estado == PORTADA and norm in _MARCAS_ABSTRACT:
            estado = ABSTRACT
            continue

        if norm in _MARCAS_REFS:
            estado = REFERENCIAS
            continue

        # Un apéndice puede comenzar con cualquier variante de la marca.
        if any(norm.startswith(m) for m in _MARCAS_APENDICE):
            estado = APENDICES
            apendices.append(t)   # el título del apéndice se incluye
            continue

        # Si estábamos en el abstract y aparece un encabezado H1,
        # el cuerpo ya comenzó (el abstract no usa estilos de encabezado APA).
        if estado == ABSTRACT and nivel == 1:
            estado = CUERPO

        # ── Acumulación por estado ────────────────────────────────────────────

        if estado == PORTADA:
            portada.append(t)

        elif estado == ABSTRACT:
            abstract.append(t)

        elif estado == CUERPO:
            # Registrar encabezados estructurados
            if nivel:
                encabezados.append(Encabezado(nivel=nivel, texto=t))
            # Detección lateral de tablas y figuras (no cambia estado)
            if _es_titulo_tabla(t):
                tablas.append(t)
            elif _es_titulo_figura(t):
                figuras.append(t)
            cuerpo.append(t)

        elif estado == REFERENCIAS:
            refs.append(t)

        elif estado == APENDICES:
            apendices.append(t)

    return DocumentoAPA(
        portada="\n".join(portada),
        abstract="\n".join(abstract),
        cuerpo="\n".join(cuerpo),
        encabezados=encabezados,
        referencias="\n".join(refs),
        apendices="\n".join(apendices),
        tablas=tablas,
        figuras=figuras,
    )


# ── API pública ───────────────────────────────────────────────────────────────

def extraer_secciones(file_bytes: bytes, completo: bool = False) -> DocumentoAPA:
    """
    Extrae las secciones de un documento APA 7 en formato .docx.

    Args:
        file_bytes: Contenido binario del archivo .docx.
        completo:   Controlado por features.extractor_completo.
                    False → modo básico (cuerpo + referencias).
                    True  → modo completo (todas las secciones APA 7).

    Returns:
        DocumentoAPA con las secciones detectadas.
        Los campos vacíos simplemente no fueron encontrados en el documento.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    return _extraer_completo(doc) if completo else _extraer_basico(doc)
