"""
Fase 0.3 — Validación de Formato Físico APA 7
=============================================
Inspecciona el documento con python-docx (sin LLM) y detecta
desviaciones del estándar APA 7ª edición.

Reglas implementadas:
  F01 — Fuente: Times New Roman 12pt | Calibri 11pt | Arial 11pt |
                Lucida Sans Unicode 10pt | Georgia 11pt  (Sección 2.19)
  F02 — Márgenes: 2.54 cm (1 in) en todos los lados            (Sección 2.17)
  F03 — Interlineado: doble en párrafos de cuerpo               (Sección 2.21)
  F04 — Sangría primera línea: 1.27 cm (0.5 in) en cuerpo       (Sección 2.23)
  F05 — Tamaño de página: carta (21.59×27.94 cm) o A4            (Sección 2.17)
  F06 — Numeración de páginas presente en encabezado             (Sección 2.18)

Todas las reglas son determinísticas; ninguna llama al LLM.
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn


# ── Constantes APA 7 ──────────────────────────────────────────────────────────

_EMU_INCH  = 914_400          # 1 pulgada en EMUs
_EMU_CM    = 360_000          # 1 cm en EMUs
_TOL_MARG  = int(0.1 * _EMU_CM)   # ±1 mm para márgenes
_TOL_PG    = int(0.5 * _EMU_CM)   # ±5 mm para tamaño de página
_TOL_IND   = int(0.1 * _EMU_INCH) # ±0.1 in para sangría

# (nombre_fuente, tamaño_pt) — APA 7, Sección 2.19
_FUENTES_VALIDAS: set[Tuple[str, int]] = {
    ("Times New Roman", 12),
    ("Calibri",         11),
    ("Arial",           11),
    ("Lucida Sans Unicode", 10),
    ("Georgia",         11),
}

# Tamaños de página en EMUs
_CARTA = (7_772_400, 10_058_400)   # 8.5 × 11 in
_A4    = (7_560_000, 10_692_000)   # 21.0 × 29.7 cm

# Sangría esperada: 0.5 in = 457 200 EMUs
_SANGRIA = int(0.5 * _EMU_INCH)

# Estilos de Word que NO son párrafos de cuerpo
_ESTILOS_NO_CUERPO = {
    "heading 1", "heading 2", "heading 3", "heading 4", "heading 5",
    "title", "subtitle", "caption",
}


# ── Modelo de error ───────────────────────────────────────────────────────────

@dataclass
class ErrorFormato:
    """Un error de formato físico detectado mediante python-docx."""
    codigo:           str   # F01 … F06
    descripcion:      str
    valor_encontrado: str
    valor_esperado:   str
    severidad:        str   # "alta" | "media" | "baja"
    regla_apa:        str
    sugerencia:       str

    def to_dict(self) -> dict:
        return {
            "codigo":           self.codigo,
            "descripcion":      self.descripcion,
            "valor_encontrado": self.valor_encontrado,
            "valor_esperado":   self.valor_esperado,
            "severidad":        self.severidad,
            "regla_apa":        self.regla_apa,
            "sugerencia":       self.sugerencia,
        }


# ── Punto de entrada público ──────────────────────────────────────────────────

def validar_formato(docx_bytes: bytes) -> List[ErrorFormato]:
    """
    Recibe los bytes de un .docx y retorna la lista de ErrorFormato.
    Lista vacía indica que el documento cumple el formato APA 7.
    """
    doc = Document(BytesIO(docx_bytes))
    errores: List[ErrorFormato] = []

    _verificar_tamano_pagina(doc, errores)
    _verificar_margenes(doc, errores)
    _verificar_fuente(doc, errores)
    _verificar_interlineado(doc, errores)
    _verificar_sangria(doc, errores)
    _verificar_numeracion_paginas(doc, errores)

    return errores


# ── Helpers ───────────────────────────────────────────────────────────────────

def _es_parrafo_cuerpo(p) -> bool:
    """True si el párrafo es texto de cuerpo (no encabezado, no vacío)."""
    if not p.text.strip():
        return False
    estilo = (p.style.name or "").lower()
    return not any(ne in estilo for ne in _ESTILOS_NO_CUERPO)


def _fuente_de_run(run) -> Tuple[Optional[str], Optional[int]]:
    """
    Retorna (nombre, tamaño_pt) resolviendo la cadena de herencia de estilos.
    Devuelve (None, None) cuando no se puede determinar.
    """
    nombre = run.font.name
    tamano = run.font.size   # Emu | None  (1 pt = 12 700 EMUs)

    # Intentar resolver desde el estilo del run
    if nombre is None or tamano is None:
        try:
            rf = run.style.font
            nombre = nombre or rf.name
            tamano = tamano or rf.size
        except Exception:
            pass

    tamano_pt = int(tamano / 12_700) if tamano else None
    return nombre, tamano_pt


def _resolver_line_spacing(p) -> Tuple[Optional[float], Optional[object]]:
    """
    Recorre la cadena de herencia de estilos para obtener
    (line_spacing, line_spacing_rule) del párrafo.
    """
    pf = p.paragraph_format
    spacing, rule = pf.line_spacing, pf.line_spacing_rule

    if spacing is not None or rule is not None:
        return spacing, rule

    estilo = p.style
    while estilo is not None:
        pf = estilo.paragraph_format
        s, r = pf.line_spacing, pf.line_spacing_rule
        if s is not None or r is not None:
            return s, r
        estilo = estilo.base_style

    return None, None


def _es_doble_espaciado(p) -> Optional[bool]:
    """
    Retorna True si doble, False si otro valor explícito, None si no determinado.
    """
    spacing, rule = _resolver_line_spacing(p)

    if rule == WD_LINE_SPACING.DOUBLE:
        return True
    if rule == WD_LINE_SPACING.MULTIPLE and spacing is not None:
        return abs(float(spacing) - 2.0) < 0.05
    if rule in (WD_LINE_SPACING.SINGLE, WD_LINE_SPACING.ONE_POINT_FIVE):
        return False
    if rule in (WD_LINE_SPACING.AT_LEAST, WD_LINE_SPACING.EXACTLY) and spacing is not None:
        return False

    return None   # herencia no resuelta — no concluyente


def _resolver_sangria(p) -> Optional[int]:
    """Retorna la sangría de primera línea en EMUs, o None si no determinada."""
    indent = p.paragraph_format.first_line_indent

    if indent is not None:
        return int(indent)

    estilo = p.style
    while estilo is not None:
        indent = estilo.paragraph_format.first_line_indent
        if indent is not None:
            return int(indent)
        estilo = estilo.base_style

    return None


# ── Verificaciones ────────────────────────────────────────────────────────────

def _verificar_tamano_pagina(doc: Document, errores: List[ErrorFormato]) -> None:
    """F05 — Tamaño de página: carta o A4."""
    for seccion in doc.sections:
        w, h = seccion.page_width, seccion.page_height
        carta_ok = abs(w - _CARTA[0]) < _TOL_PG and abs(h - _CARTA[1]) < _TOL_PG
        a4_ok    = abs(w - _A4[0])    < _TOL_PG and abs(h - _A4[1])    < _TOL_PG

        if not carta_ok and not a4_ok:
            errores.append(ErrorFormato(
                codigo="F05",
                descripcion="Tamaño de página no es Carta ni A4",
                valor_encontrado=f"{round(w/_EMU_CM,1)} × {round(h/_EMU_CM,1)} cm",
                valor_esperado="Carta (21.59 × 27.94 cm) o A4 (21.0 × 29.7 cm)",
                severidad="media",
                regla_apa="APA 7, Sección 2.17",
                sugerencia=(
                    "Ve a Diseño > Tamaño de página y selecciona "
                    "Carta o A4."
                ),
            ))
            break   # una sola vez por documento es suficiente


def _verificar_margenes(doc: Document, errores: List[ErrorFormato]) -> None:
    """F02 — Márgenes de 1 pulgada (2.54 cm) en todos los lados."""
    lados = {
        "superior":  "top_margin",
        "inferior":  "bottom_margin",
        "izquierdo": "left_margin",
        "derecho":   "right_margin",
    }
    reportados: set[str] = set()

    for seccion in doc.sections:
        for nombre, attr in lados.items():
            if nombre in reportados:
                continue
            valor = getattr(seccion, attr)
            if valor is None:
                continue
            if abs(int(valor) - _EMU_INCH) > _TOL_MARG:
                reportados.add(nombre)
                errores.append(ErrorFormato(
                    codigo="F02",
                    descripcion=f"Margen {nombre} incorrecto",
                    valor_encontrado=f"{round(int(valor)/_EMU_CM, 2)} cm",
                    valor_esperado="2.54 cm (1 pulgada)",
                    severidad="alta",
                    regla_apa="APA 7, Sección 2.17",
                    sugerencia=(
                        f"Ajusta el margen {nombre} a 2.54 cm desde "
                        "Diseño > Márgenes > Márgenes personalizados."
                    ),
                ))


def _verificar_fuente(doc: Document, errores: List[ErrorFormato]) -> None:
    """F01 — Fuente aprobada por APA 7."""
    conteo: dict[Tuple[str, int], int] = {}

    for p in doc.paragraphs:
        if not _es_parrafo_cuerpo(p):
            continue
        for run in p.runs:
            if not run.text.strip():
                continue
            nombre, tamano = _fuente_de_run(run)
            if nombre and tamano:
                key = (nombre, tamano)
                conteo[key] = conteo.get(key, 0) + len(run.text)

    if not conteo:
        return

    total = sum(conteo.values())
    dominante, chars = max(conteo.items(), key=lambda x: x[1])

    if dominante not in _FUENTES_VALIDAS and (chars / total) > 0.30:
        errores.append(ErrorFormato(
            codigo="F01",
            descripcion="Fuente no aprobada por APA 7",
            valor_encontrado=f"{dominante[0]} {dominante[1]}pt",
            valor_esperado=(
                "Times New Roman 12pt | Calibri 11pt | Arial 11pt | "
                "Lucida Sans Unicode 10pt | Georgia 11pt"
            ),
            severidad="alta",
            regla_apa="APA 7, Sección 2.19",
            sugerencia=(
                f"Reemplaza '{dominante[0]} {dominante[1]}pt' por una fuente "
                "aprobada por APA 7. La más usada en trabajos académicos es "
                "Times New Roman 12pt."
            ),
        ))


def _verificar_interlineado(doc: Document, errores: List[ErrorFormato]) -> None:
    """F03 — Interlineado doble en párrafos de cuerpo."""
    parrafos = [p for p in doc.paragraphs if _es_parrafo_cuerpo(p)]
    if not parrafos:
        return

    no_doble = [p for p in parrafos if _es_doble_espaciado(p) is False]
    proporcion = len(no_doble) / len(parrafos)

    if proporcion > 0.40:
        errores.append(ErrorFormato(
            codigo="F03",
            descripcion="Interlineado no es doble en una parte significativa del documento",
            valor_encontrado=(
                f"{len(no_doble)} de {len(parrafos)} párrafos "
                "con interlineado distinto al doble"
            ),
            valor_esperado="Interlineado doble (2.0) en todo el documento",
            severidad="alta",
            regla_apa="APA 7, Sección 2.21",
            sugerencia=(
                "Selecciona todo el texto (Ctrl+A) y aplica interlineado doble: "
                "Inicio > Párrafo > Interlineado > Doble."
            ),
        ))


def _verificar_sangria(doc: Document, errores: List[ErrorFormato]) -> None:
    """F04 — Sangría de primera línea de 0.5 in (1.27 cm)."""
    parrafos = [p for p in doc.paragraphs if _es_parrafo_cuerpo(p)]
    if not parrafos:
        return

    sin_sangria = []
    for p in parrafos:
        ind = _resolver_sangria(p)
        if ind is None:
            sin_sangria.append(p)
        elif abs(ind - _SANGRIA) > _TOL_IND:
            sin_sangria.append(p)

    proporcion = len(sin_sangria) / len(parrafos)

    if proporcion > 0.50:
        errores.append(ErrorFormato(
            codigo="F04",
            descripcion="Sangría de primera línea incorrecta o ausente",
            valor_encontrado=(
                f"{len(sin_sangria)} de {len(parrafos)} párrafos "
                "sin sangría de primera línea correcta"
            ),
            valor_esperado="1.27 cm (0.5 pulgadas) en la primera línea de cada párrafo",
            severidad="media",
            regla_apa="APA 7, Sección 2.23",
            sugerencia=(
                "Aplica sangría de primera línea de 1.27 cm a todos los párrafos: "
                "Párrafo > Sangría > Especial > Primera línea: 1.27 cm."
            ),
        ))


def _verificar_numeracion_paginas(doc: Document, errores: List[ErrorFormato]) -> None:
    """F06 — Número de página en el encabezado del documento."""
    tiene_numero = False

    for seccion in doc.sections:
        for header in (seccion.header, seccion.first_page_header):
            if header is None:
                continue
            for p in header.paragraphs:
                # Busca campos PAGE de Word en el XML del párrafo
                for elem in p._p.iter():
                    tag = elem.tag
                    if tag == qn("w:instrText") and "PAGE" in (elem.text or "").upper():
                        tiene_numero = True
                        break
                    if tag == qn("w:fldChar"):
                        # Puede ser un campo simple; marcar y verificar instrText cercano
                        pass
                if tiene_numero:
                    break
            if tiene_numero:
                break
        if tiene_numero:
            break

    if not tiene_numero:
        errores.append(ErrorFormato(
            codigo="F06",
            descripcion="Numeración de páginas no detectada en el encabezado",
            valor_encontrado="Sin campo de número de página en el encabezado",
            valor_esperado=(
                "Número de página en la esquina superior derecha, "
                "comenzando desde la página de título"
            ),
            severidad="alta",
            regla_apa="APA 7, Sección 2.18",
            sugerencia=(
                "Inserta el número de página desde "
                "Insertar > Número de página > Parte superior de la página > "
                "Número sin formato 3 (alineado a la derecha)."
            ),
        ))
