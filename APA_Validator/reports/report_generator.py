"""
Generador de reportes Word — Fase 11.1
========================================
Produce un documento .docx con el feedback del análisis APA 7,
con la identidad visual de la universidad institucional.
"""

import io

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.branding import BrandingUniversidad


def generar_reporte_docx(
    feedback:       str,
    nombre_archivo: str,
    branding:       BrandingUniversidad | None = None,
) -> io.BytesIO:
    """
    Crea un documento Word con el reporte del análisis APA 7.

    Args:
        feedback:       Texto Markdown del análisis (feedback_texto del AnalisisAPA).
        nombre_archivo: Nombre original del .docx analizado.
        branding:       Identidad visual de la universidad. Si es None, usa defaults.

    Returns:
        BytesIO con el contenido del .docx listo para descargar.
    """
    marca = branding or BrandingUniversidad()
    r, g, b = marca.color_primario_rgb()
    color = RGBColor(r, g, b)

    doc = Document()

    # ── Encabezado institucional ──────────────────────────────────────────────

    # Logo (si existe localmente)
    if marca.logo_existe:
        parrafo_logo = doc.add_paragraph()
        parrafo_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = parrafo_logo.add_run()
        run_logo.add_picture(marca.logo_path, width=Inches(1.8))

    # Nombre de la universidad
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_univ = p_univ.add_run(marca.nombre)
    run_univ.bold = True
    run_univ.font.size = Pt(13)
    run_univ.font.color.rgb = color

    # Título del reporte
    titulo = doc.add_heading("Reporte de Revisión APA 7", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = color

    # Nombre del archivo analizado
    p_archivo = doc.add_paragraph()
    p_archivo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_archivo = p_archivo.add_run(nombre_archivo)
    run_archivo.italic = True
    run_archivo.font.size = Pt(10)
    run_archivo.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()   # espacio

    # ── Contenido del feedback ────────────────────────────────────────────────

    for linea in feedback.split("\n"):
        linea = linea.strip()
        if not linea:
            continue

        # Encabezados Markdown (##)
        if linea.startswith("## "):
            h = doc.add_heading(linea[3:], level=2)
            for run in h.runs:
                run.font.color.rgb = color

        elif linea.startswith("# "):
            h = doc.add_heading(linea[2:], level=1)
            for run in h.runs:
                run.font.color.rgb = color

        # Listas con guion o asterisco
        elif linea.startswith(("- ", "* ")):
            p = doc.add_paragraph(linea[2:], style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)

        else:
            doc.add_paragraph(linea)

    # ── Pie de página ─────────────────────────────────────────────────────────

    doc.add_paragraph()
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = p_pie.add_run("Generado por Agente APA 7 · Plataforma de Revisión Académica")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0xA7, 0xA8, 0xAA)   # Gris Cool Gray 6

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
